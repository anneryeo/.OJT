from docx import Document


DOCX_PATH = r"D:\Developer\Projects\.OJT\.local\REYES_MONTHLY_OJT_MONITORING_FORM_UPDATED_WORKING.docx"


def set_cell(cell, text):
    cell.text = ""
    parts = text.split("\n\n")
    for index, part in enumerate(parts):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.text = part


doc = Document(DOCX_PATH)

set_cell(
    doc.tables[1].cell(0, 0),
    (
        "Objective 1: Organize and structure student portfolio information into a consistent data model that can support search, filtering, publication status, and future relational database integration.\n\n"
        "Objective 2: Design and implement a professional Student Portfolio Dashboard aligned with Mapua University and SOIT presentation needs, including a public directory, detailed student profile view, visitor submission form, and responsive institutional interface.\n\n"
        "Objective 3: Develop a private administration workflow for reviewing, editing, approving, returning with comments, archiving, restoring, and deleting student profile records while preserving session state in the prototype.\n\n"
        "Objective 4: Prepare defense-ready documentation, including product requirements, system architecture, user flow, database schema, running instructions, presentation structure, and updated project-document content."
    ),
)

activities = [
    (
        "May 22 - May 23",
        "Defined the scope of the Student Portfolio Dashboard, reviewed the needs of the SOIT portfolio workflow, and identified the major user groups for the system.",
        "20 hours",
        "Project scope, user roles, initial dashboard requirements",
        "Requirements Analysis, Academic Project Planning",
    ),
    (
        "May 25 - May 29",
        "Structured the student profile data model, identified academic and portfolio fields, and prepared the relational database direction for future production use.",
        "50 hours",
        "Structured content model and initial database schema direction",
        "Data Modeling, SQL Schema Planning",
    ),
    (
        "June 1 - June 5",
        "Cleaned and standardized sample student records, including year level, course category, availability, skills, profile biography, selected outcomes, and project evidence.",
        "50 hours",
        "Cleaned profile dataset with more than thirty sample records",
        "Data Cleaning, Data Structuring, Information Architecture",
    ),
    (
        "June 8 - June 12",
        "Designed and implemented the public-facing dashboard interface using an institutional red, gold, and white visual direction appropriate for a Mapua University prototype.",
        "50 hours",
        "Public directory layout, profile cards, hero section, responsive styling",
        "UI/UX Design, Frontend Development, Accessibility Awareness",
    ),
    (
        "June 15 - June 19",
        "Integrated the structured student data into the static frontend, implemented search, filters, sorting, detailed profile panels, and the public profile submission form.",
        "50 hours",
        "Functional public portfolio dashboard and visitor submission workflow",
        "JavaScript Rendering, State Handling, Interface Logic",
    ),
    (
        "June 22 - June 26",
        "Developed the private admin route with fake login, session persistence, admin editor, submission queue, published records, archive/disapproval records, and moderation actions.",
        "50 hours",
        "Private admin CMS prototype with add, edit, save, approve, return, archive, restore, and delete functions",
        "Workflow Design, Local Storage Persistence, QA Testing",
    ),
    (
        "June 29 - July 3",
        "Finalized the defense-ready documentation, SQL schema, user flow, architecture notes, presentation structure, project-document updates, and final quality assurance checks.",
        "50 hours",
        "Updated PRD, database schema, presentation structure, project document copy, and rendered DOCX QA outputs",
        "Technical Writing, Documentation, Defense Preparation",
    ),
    (
        "July 4",
        "Reviewed the completed prototype and documentation package to ensure alignment with the Chapter 3 project overview, objectives, significance, and limitations.",
        "20 hours",
        "Final review notes and defense-ready prototype package",
        "Project Evaluation, Communication, Final QA",
    ),
]

for row_index, row_data in enumerate(activities, start=1):
    for col_index, value in enumerate(row_data):
        doc.tables[2].cell(row_index, col_index).text = value

set_cell(
    doc.tables[3].cell(0, 0),
    (
        "During this internship period, I learned how to translate unstructured student portfolio information into a coherent data model and a working dashboard prototype. The project strengthened my understanding that a visually effective dashboard depends first on consistent data structure, clear field definitions, and a workflow that protects data quality before publication.\n\n"
        "I also gained practical experience in building an end-to-end prototype using HTML, CSS, and JavaScript. The final system includes a public portfolio directory, searchable and filterable student cards, expanded profile panels, a visitor submission form, and a private admin route. Through the admin workflow, I learned how moderation states such as pending, returned, published, archived, restored, and deleted records affect both user experience and data governance.\n\n"
        "The documentation phase also became a major learning point. I prepared product requirements, a database schema, a user flow, a system architecture explanation, running instructions, and a presentation outline. This process helped me connect technical implementation with academic reporting and defense preparation."
    ),
)

set_cell(
    doc.tables[4].cell(0, 0),
    (
        "One of the main challenges was aligning the prototype with the actual project scope while avoiding features that would require external systems or live university integrations. The system needed to be realistic enough for a production path, but still feasible as a static prototype within the OJT period. I addressed this by using browser local storage for session persistence and CMS state while documenting how the same workflow could later connect to a relational database and authenticated backend.\n\n"
        "Another challenge was ensuring that the admin CMS felt realistic without making it publicly visible. The solution was to separate the public dashboard from the private admin route, add a fake login for demonstration, and provide admin features such as add, edit, save, approve, return with comments, archive, restore, and delete. I also refined the public submission form to remove redundant fields and infer course category from the program input.\n\n"
        "A further challenge involved keeping the documentation consistent with the evolving prototype. I resolved this by updating the local documentation, project document copy, and presentation structure after the interface and workflow changes were finalized."
    ),
)

set_cell(
    doc.tables[5].cell(0, 0),
    (
        "The feedback received during the final refinement stage emphasized that the prototype should reflect a more realistic academic and administrative workflow. In response, the admin CMS was moved out of the public scroll flow and placed under a private admin route. The profile submission form was simplified by removing the redundant course-type dropdown, while the admin workspace was expanded to support editing, saving, returning with comments, archiving, restoring, deleting, searching, filtering, and limiting record lists to a manageable top-five view.\n\n"
        "The feedback also reinforced the need for the documentation to match the implemented system rather than an earlier concept. As a result, the PRD, database schema, defense README, presentation structure, and project document copy were revised to describe the actual outcome: a static but production-shaped Student Portfolio Dashboard with a documented path toward a real database-backed CMS."
    ),
)

self_assessment_rows = {
    1: (
        "Task Completion",
        "5",
        "Completed the public portfolio dashboard, private admin route, profile submission workflow, moderation tools, updated database schema, and defense documentation package.",
    ),
    2: (
        "Initiative & Proactivity",
        "5",
        "Independently refined the prototype beyond the initial directory concept by adding a realistic CMS workflow, private admin access, local persistence, and production-readiness documentation.",
    ),
    3: (
        "Communication",
        "4",
        "Documented the system clearly through the PRD, README, database schema, presentation outline, and project document updates so that the prototype can be explained during defense.",
    ),
    4: (
        "Time Management",
        "5",
        "Managed the compressed 320-hour deployment by sequencing implementation, QA, documentation, and document revision tasks within the required period.",
    ),
    5: (
        "Learning Engagement",
        "5",
        "Applied new learning in data modeling, frontend state management, administrative workflow design, technical documentation, and DOCX render verification.",
    ),
}

for row_index, values in self_assessment_rows.items():
    for col_index, value in enumerate(values):
        doc.tables[6].cell(row_index, col_index).text = value

set_cell(
    doc.tables[7].cell(0, 0),
    (
        "Christine demonstrated strong ownership of the Student Portfolio Dashboard prototype by progressing from data organization and interface design to a more complete administrative workflow. The final output reflects a clearer understanding of how student portfolio records should be structured, moderated, and presented for academic review. The prototype now includes a public directory, visitor submission flow, private admin route, fake login, session persistence, profile editing, approval, return with comments, archiving, restoration, deletion, and defense-ready documentation.\n\n"
        "Her work shows initiative in aligning the implementation with the project document's objectives and limitations. The system remains appropriately scoped as a prototype, but it includes a clear production path through the documented relational database schema, future authentication requirements, validation needs, and audit-log planning."
    ),
)

doc.save(DOCX_PATH)
print(DOCX_PATH)
