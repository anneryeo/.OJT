from docx import Document
from docx.oxml.ns import qn


DOCX_PATH = r"D:\Developer\Projects\.OJT\.local\REYES_CHRISTINE_JULLIANE_OJT_Project_Document_Prototype_Updated.docx"
BODY_STYLE = "Thesis: Paragraph Text"


def set_paragraph(paragraph, text, style=None):
    paragraph.text = ""
    if style:
        paragraph.style = style
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is not None:
        ppr.remove(num_pr)
    paragraph.add_run(text)


doc = Document(DOCX_PATH)

updates = {
    74: (
        "The project is a dynamic and interactive Student Portfolio Dashboard for the School of Information Technology "
        "that centralizes student achievements, academic standing, portfolio evidence, and project history in a single "
        "accessible interface. In its current prototype version, the system is presented as the Mapua Student Portfolio "
        "Registry, a web-based dashboard that allows public visitors to discover student profiles while providing a "
        "separate private administration workspace for profile review and data governance. The prototype contains "
        "thirty-three sample student profiles and demonstrates how structured student records can be transformed into "
        "a searchable and visually readable dashboard."
    ),
    76: (
        "Prior to the proposed system, student information and portfolio artifacts were treated as scattered records "
        "rather than as a unified academic data resource. This made it difficult for faculty and administrators to "
        "assess a student's holistic academic journey, compare project evidence across program categories, or prepare "
        "student capabilities for institutional partners and accreditation-related presentations. The opportunity "
        "addressed by the project is the creation of a cleaned, structured, and moderated portfolio pipeline that "
        "turns unstructured student data into a reliable frontend dashboard supported by a production-ready relational "
        "schema."
    ),
    78: (
        "The first objective of the project is to convert unstructured student data into a dependable relational data "
        "model. The prototype demonstrates this objective through normalized student records, course classification, "
        "year-level status, portfolio evidence, skills, and moderation states. The supporting SQL schema documents how "
        "student identity, academic details, skills, projects, metrics, administrator users, and audit logs may be "
        "represented in a production database."
    ),
    79: (
        "The second objective is to design a usable frontend dashboard that supports both discovery and academic "
        "administration. The public-facing interface allows visitors to search and filter student portfolios by year "
        "level, course category, availability, and keywords. The private administration interface, reached through the "
        "prototype's admin route, allows authorized users in a future production version to maintain data quality by "
        "reviewing, editing, approving, returning, archiving, restoring, and deleting records."
    ),
    80: (
        "The third objective is to connect the data pipeline to an interface that reflects updates immediately. In the "
        "prototype, this is simulated through browser local storage, which persists submitted profiles, administrator "
        "edits, approval decisions, returned comments, archived records, and session state. Although the prototype is "
        "static, its data structures are intentionally designed so that a backend API and database can replace local "
        "storage without redesigning the user workflow."
    ),
    82: (
        "The project reduces administrative friction by transforming student profiles into structured, searchable, and "
        "moderated records. Faculty and administrators can more easily inspect student capability, monitor the quality "
        "of submitted portfolio information, and prepare a more organized view of student work for external partners. "
        "The system also benefits students by providing a clearer publication process for presenting skills, projects, "
        "and availability in a professional format aligned with Mapua University's academic and institutional identity."
    ),
    84: (
        "The prototype is scoped to portfolio data for School of Information Technology programs and related computing, "
        "data, information systems, information technology, computer science, and media or design-oriented student work. "
        "It does not perform live synchronization with external Mapua systems such as Blackboard, nor does it use real "
        "cloud authentication, email verification, file uploads, or a hosted production database. For the purpose of "
        "the 320-hour deployment, these integrations remain outside the scope, while the prototype focuses on data "
        "structuring, dashboard interaction, administrative moderation, and defense-ready documentation."
    ),
    86: (
        "The final prototype is documented through the public interface, the private administration route, and the "
        "supporting product and database documents stored in the local documentation folder. The following figures "
        "describe the major screens and workflows that demonstrate the system from data presentation to administrative "
        "governance."
    ),
    87: "Figure 3.1: Public landing section of the Mapua Student Portfolio Registry",
    88: (
        "The landing section introduces the Student Portfolio Dashboard as a professional registry for Mapua student "
        "work. It uses a school-aligned red, gold, and white visual system and communicates the purpose of the site as "
        "a centralized discovery interface for student achievements, academic status, and project evidence. The summary "
        "metrics show the number of published profiles, pending submissions, featured projects, and course categories, "
        "giving administrators and visitors a quick overview of the current portfolio dataset."
    ),
    89: (
        "This section also establishes the institutional tone of the prototype. Rather than functioning only as a "
        "decorative landing page, it immediately frames the website as a data-driven portfolio registry that can support "
        "faculty review, partner discovery, and OJT presentation requirements."
    ),
    90: "",
    91: "Figure 3.2: Public directory with search, filters, and student cards",
    92: (
        "The public directory contains the published student profiles and allows visitors to search by name, program, "
        "skills, project details, and other profile content. The directory includes filters for year level, course type, "
        "and availability, while sorting options allow records to be arranged by featured status, name, or year level. "
        "The current course-type filter focuses on Computer Science, Information Technology, Information Systems, Data "
        "Science, and Media and Design, while uncategorized or broader technology records remain visible through the "
        "All course types view."
    ),
    93: (
        "Each card is generated from structured data rather than hard-coded profile markup. This supports the project's "
        "data-cleaning objective because the same rendering logic can display seed records, approved visitor submissions, "
        "or future records loaded from a database-backed API."
    ),
    94: "",
    95: "",
    96: "Figure 3.3: Expanded student profile panel",
    97: (
        "When a visitor selects a student card, the system opens an expanded profile panel that presents the student's "
        "name, course category, program, year-level status, role, location, availability, biography, skills, selected "
        "outcomes, and featured project work. This view demonstrates how a cleaned student record can be transformed "
        "into a readable portfolio summary for faculty, recruiters, mentors, or external partners."
    ),
    98: (
        "The profile panel also includes contact and sharing actions. These actions support the project's purpose of "
        "making student capabilities easier to present and circulate while still keeping the profile structure consistent "
        "across all students."
    ),
    99: "",
    100: "Figure 3.4: Student publication request form",
    101: (
        "The public-facing submission form allows students or visitors to request profile publication by entering their "
        "name, email, program, year level, portfolio URL, skills, and short biography. To avoid redundant input, the "
        "form no longer asks the submitter to choose a separate course-type dropdown. Instead, the prototype infers the "
        "course category from the program text and sends the record into the administration queue for review."
    ),
    102: (
        "Submitted profiles are not published immediately. They are stored as pending records so that an administrator "
        "can verify, clean, classify, edit, and approve the information before it appears in the public directory. This "
        "reflects the need for data governance in an academic portfolio system."
    ),
    103: "",
    104: "",
    105: "",
    106: "",
    107: "Figure 3.5: Private administration CMS and moderation workflow",
    108: (
        "The administration CMS is no longer displayed as a public page section. It is accessed through a private "
        "prototype route, such as index.html#/admin in local static mode or /admin when hosted with an appropriate "
        "rewrite rule. The route includes a fake login for demonstration purposes, where any email and password can "
        "start an admin session. This simulates authentication without requiring a live identity provider during the "
        "prototype defense."
    ),
}

for index, text in updates.items():
    if index in {87, 91, 96, 100, 107}:
        set_paragraph(doc.paragraphs[index], text, None)
    elif text:
        set_paragraph(doc.paragraphs[index], text, BODY_STYLE)
    else:
        set_paragraph(doc.paragraphs[index], text, None)

additional_paragraphs = [
    (
        "Inside the private CMS, administrators can add new profiles, edit and save existing records, approve pending "
        "submissions, return submissions with comments, archive or disapprove records, restore archived profiles, and "
        "delete records from the prototype state. The system maintains separate administrative views for submission "
        "queues, published portfolios, and archived or disapproved records. Each category includes search and course "
        "filter controls and shows the top five records by default, with an option to see more when the list is longer."
    ),
    (
        "The private CMS is intentionally backed by browser local storage in the prototype. This provides session "
        "persistence for defense demonstration while clearly separating prototype storage from the production design. "
        "In a deployed implementation, local storage should be replaced by the relational database schema, authenticated "
        "administrator accounts, server-side validation, and audit logging."
    ),
    (
        "The system architecture separates the public interface, administrative workflow, structured data, and production "
        "database plan. The public interface is implemented in index.html and styled through styles.css, while app.js "
        "controls rendering, filtering, detail panels, visitor submission, admin routing, fake login, moderation actions, "
        "and local persistence. The production schema defines tables for students, skills, projects, metrics, admin users, "
        "and audit logs, allowing the prototype to be defended as a production-shaped system even though it remains "
        "dependency-free and locally runnable."
    ),
    (
        "For the defense, the important point is that the prototype demonstrates the complete end-to-end process. A "
        "student profile can be submitted through the public website, held as pending data, reviewed in the private "
        "admin workspace, edited for quality, approved for publication, returned with comments, archived, restored, or "
        "deleted. This workflow directly supports the Chapter 3 objective of linking a cleaned backend data pipeline to "
        "a readable dashboard that can update portfolio information in real time within the prototype environment."
    ),
]

appendices = doc.paragraphs[109]._p
parent = appendices.getparent()
for text in additional_paragraphs:
    p = doc.add_paragraph()
    p.style = BODY_STYLE
    p.add_run(text)
    parent.remove(p._p)
    parent.insert(parent.index(appendices), p._p)

doc.save(DOCX_PATH)
print(DOCX_PATH)
